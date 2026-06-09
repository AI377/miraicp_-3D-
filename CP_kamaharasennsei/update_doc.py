import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt

dst = r'\\wsl.localhost\Ubuntu\home\yusuke\CP_kamaharasennsei\miraiCP_要件定義書_v1.1.docx'
doc = Document(dst)

# ---- メンバーテーブル（Table 0）を削除 ----
member_table = doc.tables[0]
member_table._tbl.getparent().remove(member_table._tbl)

# 「プロジェクトメンバー」見出しを削除
for para in list(doc.paragraphs):
    if 'プロジェクトメンバー' in para.text:
        para._element.getparent().remove(para._element)

# ---- 実装状況セクションを末尾に追加 ----
h1 = next(s for s in doc.styles if s.name == 'Heading 1')
h2 = next(s for s in doc.styles if s.name == 'Heading 2')

doc.add_paragraph()

p = doc.add_paragraph('11. 実装状況（2026年5月26日時点）')
p.style = h1

doc.add_paragraph('以下は現時点で実装・動作確認済みのコンポーネントである。')

table = doc.add_table(rows=1, cols=4)

for i, txt in enumerate(['ファイル名', '機能', '状態', '備考']):
    cell = table.rows[0].cells[i]
    cell.text = txt
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

rows_data = [
    ('check_mediapipe.py',
     'MediaPipe Handsによる手の骨格認識\n・21点ランドマーク取得\n・左右手の識別\n・骨格の画面描画',
     '動作確認済み', 'PCカメラ1台で両手同時認識可能を確認'),
    ('gesture_engine.py',
     'G5グリップ検出\n・手を握る（グー）でGRIP ON\n・手を開く（パー）でGRIP OFF\n・状態変化時のみコンソール出力',
     '動作確認済み', 'グリップ状態を画面上に赤/緑で表示'),
]

for cols in rows_data:
    row = table.add_row()
    for i, val in enumerate(cols):
        row.cells[i].text = val
        for run in row.cells[i].paragraphs[0].runs:
            run.font.size = Pt(9.5)

doc.add_paragraph()
p2 = doc.add_paragraph('次の実装予定')
p2.style = h2

for step in [
    'server.py — WebSocketサーバー（ジェスチャーコマンドをブラウザに送信）',
    'index.html / app.js — Three.jsによる3Dオブジェクト表示・操作',
    'G1ズームイン — 両手を広げる動作でオブジェクト拡大',
    'フィジカル回転 — 指スライド・フリックで慣性付き回転',
    'ゴミ箱削除 — 選択オブジェクトを払ってゴミ箱へ',
]:
    doc.add_paragraph(step)

doc.save(dst)
print('docx saved')

from docx2pdf import convert
convert(dst, dst.replace('.docx', '.pdf'))
print('pdf saved')
